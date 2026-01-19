# src/backend/pdf_generator.py
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.colors import black, grey
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumePDFGenerator:
    def __init__(self):
        # 1. Setup Styles to match Jake Ryan's LaTeX look
        self.styles = getSampleStyleSheet()
        
        # Name: Large, Centered, Bold
        self.styles.add(ParagraphStyle(
            name='NameHeader',
            fontName='Helvetica-Bold',
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=4
        ))

        # Contact Info: Small, Centered
        self.styles.add(ParagraphStyle(
            name='ContactInfo',
            fontName='Helvetica',
            fontSize=9,
            alignment=TA_CENTER,
            spaceAfter=10
        ))

        # Section Header: Uppercase, Bold, with a line underneath (handled by Table)
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            spaceBefore=6,
            spaceAfter=2,
            textTransform='uppercase' # Simulates the uppercase look
        ))

        # Content Main: Bold (e.g., University Name)
        self.styles.add(ParagraphStyle(
            name='ItemTitle',
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT
        ))

        # Content Sub: Italic/Normal (e.g., Role, Degree)
        self.styles.add(ParagraphStyle(
            name='ItemSub',
            fontName='Helvetica-Oblique',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT
        ))

        # Right Aligned Date/Location
        self.styles.add(ParagraphStyle(
            name='RightAlign',
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_RIGHT
        ))

        # Bullet Points
        self.styles.add(ParagraphStyle(
            name='BulletPoint',
            fontName='Helvetica',
            fontSize=9.5,
            leading=12,
            leftIndent=12,
            firstLineIndent=0,
            alignment=TA_LEFT,
            spaceAfter=1
        ))

    def _create_section_header(self, title):
        """Creates a section header with a horizontal line under it."""
        # We use a table to create the bottom border effect
        p = Paragraph(title.upper(), self.styles['SectionHeader'])
        data = [[p]]
        t = Table(data, colWidths=['100%'])
        t.setStyle(TableStyle([
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('LINEBELOW', (0, 0), (-1, -1), 1, colors.black),
        ]))
        return t

    def _create_split_row(self, left_text, right_text, left_style='ItemTitle', right_style='RightAlign'):
        """Creates a row with text on the left and text on the right."""
        p_left = Paragraph(left_text, self.styles[left_style])
        p_right = Paragraph(right_text, self.styles[right_style])
        data = [[p_left, p_right]]
        t = Table(data, colWidths=['75%', '25%'])
        t.setStyle(TableStyle([
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ]))
        return t

    def generate_pdf(self, resume_data: dict):
        buffer = io.BytesIO()
        # Narrow margins like the Jake Ryan template 
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            rightMargin=0.5*inch, leftMargin=0.5*inch,
            topMargin=0.5*inch, bottomMargin=0.5*inch
        )
        story = []

        # ================= HEADER =================
        story.append(Paragraph(resume_data.get("name", "Your Name"), self.styles["NameHeader"]))
        
        # Build Contact String: Email | Phone | LinkedIn | GitHub
        parts = []
        if resume_data.get("phone"): parts.append(resume_data["phone"])
        if resume_data.get("email"): parts.append(resume_data["email"])
        if resume_data.get("linkedin"): 
            parts.append(f"linkedin.com/in/{resume_data['linkedin'].split('/')[-1]}")
        if resume_data.get("github"): 
            parts.append(f"github.com/{resume_data['github'].split('/')[-1]}")
        
        contact_text = " | ".join(parts)
        story.append(Paragraph(contact_text, self.styles["ContactInfo"]))
        story.append(Spacer(1, 10))

        # ================= SUMMARY (FIXED) =================
        # Check for 'summary' OR 'professional_summary'
        summary_text = resume_data.get("summary") or resume_data.get("professional_summary")
        if summary_text:
            story.append(self._create_section_header("Professional Summary")) #
            story.append(Spacer(1, 4))
            story.append(Paragraph(summary_text, self.styles["BulletPoint"])) # Using BulletPoint style for body text look
            story.append(Spacer(1, 8))

        # ================= EDUCATION =================
        if resume_data.get("education"):
            story.append(self._create_section_header("Education"))
            story.append(Spacer(1, 4))
            
            for edu in resume_data["education"]:
                # Row 1: University (Left) | Location (Right)
                school = edu.get("institution", "University")
                loc = edu.get("location", "") # Need to ensure location is captured in app.py
                story.append(self._create_split_row(school, loc, 'ItemTitle'))
                
                # Row 2: Degree (Left) | Date (Right)
                degree = edu.get("degree", "Degree")
                date = edu.get("year") or edu.get("period") or ""
                story.append(self._create_split_row(degree, date, 'ItemSub'))
                story.append(Spacer(1, 4))

        # ================= EXPERIENCE =================
        if resume_data.get("experience"):
            story.append(self._create_section_header("Experience"))
            story.append(Spacer(1, 4))

            for exp in resume_data["experience"]:
                # App.py sends: company, role, start_date, end_date, description
                role = exp.get("role") or exp.get("title") or "Role"
                
                # Format date: "Start - End"
                start = exp.get("start_date", "")
                end = exp.get("end_date", "")
                period = f"{start} - {end}" if start else (exp.get("period") or "")
                
                story.append(self._create_split_row(role, period, 'ItemTitle'))
                
                company = exp.get("company", "Company")
                location = exp.get("location", "")
                story.append(self._create_split_row(company, location, 'ItemSub'))
                
                # Handle description (list or string)
                desc = exp.get("description") or exp.get("responsibilities") or []
                if isinstance(desc, str):
                    # Split by newline if it's a raw string from text area
                    desc = [x.strip() for x in desc.split('\n') if x.strip()]
                
                for bullet in desc:
                    # Clean up existing bullets from user input
                    text = bullet.lstrip('-').lstrip('•').strip()
                    story.append(Paragraph(f"• {text}", self.styles["BulletPoint"]))
                
                story.append(Spacer(1, 6))
        # ================= PROJECTS =================
        if resume_data.get("projects"):
            story.append(self._create_section_header("Projects"))
            story.append(Spacer(1, 4))

            for proj in resume_data["projects"]:
                title = proj.get("title", "Project")
                duration = proj.get("duration", "")
                
                # Title Row
                story.append(self._create_split_row(title, duration, 'ItemTitle'))
                
                # Tech Stack Row
                tech_stack = proj.get("tech_stack", "")
                if isinstance(tech_stack, list): 
                    tech_stack = ", ".join(tech_stack)
                if tech_stack:
                    story.append(Paragraph(f"<b>Tech Stack:</b> <i>{tech_stack}</i>", self.styles["BulletPoint"]))

                # Description/Outcome
                desc = proj.get("description") or proj.get("outcome") or []
                if isinstance(desc, str):
                    desc = [x.strip() for x in desc.split('\n') if x.strip()]
                
                for bullet in desc:
                    text = bullet.lstrip('-').lstrip('•').strip()
                    story.append(Paragraph(f"• {text}", self.styles["BulletPoint"]))
                
                story.append(Spacer(1, 6))
        # ================= SKILLS =================
        # [cite: 42]
        if resume_data.get("skills"):
            story.append(self._create_section_header(" Skills"))
            story.append(Spacer(1, 4))
            
            skills = resume_data["skills"]
            # Handle format {"technical": [], "soft": []} or list
            
            skill_lines = []
            if isinstance(skills, dict):
                if skills.get("technical"):
                    # We want: "Languages: Java, Python..."
                    # Since we don't have categories in the simple input, we list them all.
                    # Or if you have categories in your schema, use them.
                    skill_lines.append(f"<b>Technical:</b> {', '.join(skills['technical'])}")
                if skills.get("soft"):
                    skill_lines.append(f"<b>Soft Skills:</b> {', '.join(skills['soft'])}")
            elif isinstance(skills, list):
                skill_lines.append(f"<b>Skills:</b> {', '.join(skills)}")
            
            for line in skill_lines:
                story.append(Paragraph(line, self.styles["BulletPoint"]))
        # ================= CERTIFICATIONS (FIXED) =================

        if resume_data.get("certifications"):
            story.append(self._create_section_header("Certifications"))
            story.append(Spacer(1, 4))
            
            certs = resume_data["certifications"]
            # Handle if it's a string or list
            if isinstance(certs, str):
                certs = [x.strip() for x in certs.split('\n') if x.strip()]

            for cert in certs:
    # ✅ Case 1: certification is a dict (AI output)
                if isinstance(cert, dict):
                    name = cert.get("name", "").strip()
                    authority = cert.get("issuing_authority", "").strip()
                    issue_date = cert.get("issue_date", "").strip()

                    parts = [name]
                    if authority:
                        parts.append(authority)
                    if issue_date:
                        parts.append(issue_date)

                    text = " – ".join(parts)

                # ✅ Case 2: certification is a string (manual input)
                else:
                    text = str(cert).lstrip('-').lstrip('•').strip()

                if text:
                    story.append(Paragraph(f"• {text}", self.styles["BulletPoint"]))
    
            
        doc.build(story)
        buffer.seek(0)
        return buffer

        doc.build(story)
        buffer.seek(0)
        return buffer
# --- Add this import at the top of pdf_generator.py ---
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Add this class at the bottom of pdf_generator.py ---
class ResumeDOCXGenerator:
    def generate_docx(self, resume_data: dict):
        doc = Document()
        
        # 1. Name & Contact (Center Aligned)
        name = doc.add_heading(resume_data.get("name", "Your Name"), 0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_info = []
        if resume_data.get("email"): contact_info.append(resume_data["email"])
        if resume_data.get("phone"): contact_info.append(resume_data["phone"])
        if resume_data.get("linkedin"): contact_info.append(resume_data["linkedin"])
        
        contact_p = doc.add_paragraph(" | ".join(contact_info))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. Professional Summary
        if resume_data.get("professional_summary"):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(resume_data["professional_summary"])
        
        # 3. Education
        if resume_data.get("education"):
            doc.add_heading('Education', level=1)
            for edu in resume_data["education"]:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', 'Degree')}").bold = True
                p.add_run(f", {edu.get('institution', 'University')}")
                if edu.get('period'):
                    p.add_run(f" ({edu['period']})")

        # 4. Experience
        if resume_data.get("experience"):
            doc.add_heading('Experience', level=1)
            for exp in resume_data["experience"]:
                # Role at Company
                # Role (bold)
                p_role = doc.add_paragraph()
                p_role.add_run(exp.get("role", "")).bold = True

                # Company (italic)
                p_company = doc.add_paragraph()
                p_company.add_run(exp.get("company", "")).italic = True

                # Period (italic, already normalized)
                if exp.get("period"):
                    p_period = doc.add_paragraph()
                    p_period.add_run(exp["period"]).italic = True

                
                # Description
                desc = exp.get('description', [])
                if isinstance(desc, list):
                    for bullet in desc:
                        doc.add_paragraph(bullet, style='List Bullet')
                else:
                    doc.add_paragraph(desc)
        # 5. Projects
        if resume_data.get("projects"):
            doc.add_heading('Projects', level=1)

            for proj in resume_data["projects"]:
                # Title
                p_title = doc.add_paragraph()
                p_title.add_run(proj.get("title", "Project")).bold = True

                # Duration
                if proj.get("duration"):
                    p_dur = doc.add_paragraph()
                    p_dur.add_run(proj["duration"]).italic = True

                # Tech stack
                tech = proj.get("tech_stack", [])
                if isinstance(tech, list):
                    tech = ", ".join(tech)
                if tech:
                    doc.add_paragraph(f"Tech Stack: {tech}")

                # Description bullets
                desc = proj.get("description", [])
                if isinstance(desc, list):
                    for bullet in desc:
                        doc.add_paragraph(bullet, style="List Bullet")

        # 6. Skills
        skills = resume_data.get("skills")
        if skills:
            doc.add_heading("Skills", level=1)
            if isinstance(skills, dict):
                if skills.get("technical"):
                    doc.add_paragraph(
                        ", ".join(skills["technical"]),
                        style="List Bullet"
                    )

                if skills.get("soft"):
                    doc.add_paragraph(
                        ", ".join(skills["soft"]),
                        style="List Bullet"
                    )

        # 7. Certification
        certs = resume_data.get("certifications")

        if certs:
            doc.add_heading("Certifications", level=1)

            for cert in certs:
                # Case 1: certification is a dict (AI output)
                if isinstance(cert, dict):
                    name = cert.get("name", "").strip()
                    authority = cert.get("issuing_authority", "").strip()
                    year = cert.get("issue_date", "").strip()

                    parts = [name]
                    if authority:
                        parts.append(authority)
                    if year:
                        parts.append(year)

                    text = " – ".join(parts)

                # Case 2: certification is a string (manual input)
                else:
                    text = str(cert).strip()

                if text:
                    doc.add_paragraph(text, style="List Bullet")


        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer